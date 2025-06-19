import os
import base64
from pathlib import Path
import streamlit as st
import pypandoc
import PyPDF2
import difflib
import docx
from docx import Document
import fitz  # PyMuPDF
from collections import defaultdict
import re

from src.preprocessing.text_extract import docx_to_text, pdf_to_text
from src.preprocessing.scan_text_extract import image_to_text
from src.preprocessing.pdf_to_image import pdf_to_images

from difflib import HtmlDiff

# def generate_diff_html(text1: str, text2: str) -> str:
#     lines1 = text1.splitlines()
#     lines2 = text2.splitlines()
#     diff = HtmlDiff().make_table(lines1, lines2, fromdesc="Document 1", todesc="Document 2")
#     return diff


def rotate_image_to_portrait(image):
    """Rotate image to portrait orientation if needed"""
    width, height = image.size
    if width > height:
        return image.rotate(270, expand=True)
    return image

def extract_text_from_file(file_path, file_type):
    """Extrait le texte d'un fichier selon son type"""
    try:
        print(f"Tentative d'extraction du texte de: {file_path}")

        if file_type in ['.jpg', '.jpeg', '.png']:
            text = image_to_text(file_path)
            print(f"Texte extrait de l'image (longueur: {len(text)}): {text[:100]}...")
            return text

        elif file_type == '.pdf':
            text = pdf_to_text(file_path)
            print(f"Texte extrait du PDF (longueur: {len(text)}): {text[:100]}...")

            if not text.strip() or len(text.strip()) < 50:
                print("PDF probablement scanné, conversion en images...")
                image_paths = pdf_to_images(file_path)
                if not image_paths:
                    raise ValueError("Impossible de convertir le PDF en images")

                all_text = []
                for img_path in image_paths:
                    print(f"Traitement de l'image: {img_path}")
                    page_text = image_to_text(img_path)
                    print(f"Texte extrait de la page (longueur: {len(page_text)}): {page_text[:100]}...")
                    if page_text:
                        all_text.append(page_text)

                for img_path in image_paths:
                    if os.path.exists(img_path):
                        os.remove(img_path)

                final_text = "\n".join(all_text)
                print(f"Texte final du PDF scanné (longueur: {len(final_text)}): {final_text[:100]}...")
                return final_text
            return text

        elif file_type == '.docx':
            text = docx_to_text(file_path)
            print(f"Texte extrait du DOCX (longueur: {len(text)}): {text[:100]}...")
            return text

        else:
            raise ValueError(f"Type de fichier non supporté: {file_type}")

    except Exception as e:
        print(f"Erreur lors de l'extraction du texte: {str(e)}")
        return ""

def display_pdf(file_path):
    """Affiche un fichier PDF dans Streamlit via une iframe"""
    with open(file_path, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode("utf-8")
    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="600px" type="application/pdf"></iframe>'
    st.markdown(pdf_display, unsafe_allow_html=True)

def convert_docx_to_pdf(docx_path, output_dir="temp"):
    """Convertit un fichier DOCX en PDF avec Pandoc"""
    pdf_path = Path(output_dir) / (Path(docx_path).stem + ".pdf")
    try:
        pypandoc.convert_file(str(docx_path), 'pdf', outputfile=str(pdf_path))
        return str(pdf_path)
    except Exception as e:
        print(f"Erreur lors de la conversion DOCX → PDF : {e}")
        return None

def count_pages(file_path, file_type):
    """Compte le nombre de pages d'un document"""
    try:
        if file_type in ['.jpg', '.jpeg', '.png']:
            return 1
        elif file_type == '.pdf':
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                return len(reader.pages)
        elif file_type == '.docx':
            doc = Document(file_path)
            # Estimation basée sur le nombre de ruptures de section
            return len(doc.sections)
        else:
            raise ValueError(f"Type de fichier non supporté: {file_type}")
    except Exception as e:
        print(f"Erreur lors de la comptage des pages: {str(e)}")
        return 0

def segment_text_by_topics(text):
    """Segmente le texte en sections en fonction de modèles de numérotation"""
    import re
    
    # Modèle pour identifier les sections (ex: 1., 1.1., 2., etc)
    topic_pattern = r'^\s*(\d+(?:\.\d+)*\.?)\s+(.+)$'
    
    # Diviser le texte en lignes
    lines = text.split('\n')
    
    # Dictionnaire pour stocker les sections
    topics = {}
    current_topic = None
    current_content = []
    
    for line in lines:
        match = re.match(topic_pattern, line)
        if match:
            # Si nous avons une section actuel, la sauvegarder
            if current_topic:
                topics[current_topic] = '\n'.join(current_content)
            
            # Commencer un nouveau sujet
            current_topic = match.group(1)
            current_content = [match.group(2)]
        elif current_topic:
            current_content.append(line)
    
    # Ajouter la dernière section
    if current_topic:
        topics[current_topic] = '\n'.join(current_content)
    
    return topics




def analyze_text_structure(pdf_path):
    doc = fitz.open(pdf_path)
    size_stats = defaultdict(list)

    for page in doc:
        for block in page.get_text("dict")["blocks"]:
            for line in block.get("lines", []):
                for span in line["spans"]:
                    size = round(span["size"], 1)
                    text = span["text"].strip()
                    if text:
                        size_stats[size].append(text)

    # Trie les tailles de police décroissantes (hiérarchie)
    distribution = sorted(size_stats.items(), key=lambda x: -x[0])

    # Classement hiérarchique
    hierarchy = {}
    for i, (size, texts) in enumerate(distribution):
        if i == 0:
            label = "Grand titre"
        elif i == 1:
            label = "Titre"
        elif i == 2:
            label = "Sous-titre"
        else:
            label = f"Paragraphe niveau {i - 2}"
        hierarchy[label] = {
            "taille_px": size,
            "occurrences": len(texts),
            "exemples": texts[:5]
        }

    return hierarchy


def extract_typo_blocks(pdf_path):
    """
    Extrait les blocs de texte d’un PDF avec leur taille de police et leur position.
    Retourne une liste de dicts : [{text, size, page}, ...]
    """
    import fitz  # PyMuPDF
    blocks = []

    doc = fitz.open(pdf_path)
    for page_number, page in enumerate(doc, start=1):
        for block in page.get_text("dict")["blocks"]:
            for line in block.get("lines", []):
                for span in line["spans"]:
                    text = span.get("text", "").strip()
                    if text:
                        blocks.append({
                            "text": text,
                            "size": round(span["size"], 1),
                            "page": page_number
                        })
    return blocks


# def extract_paragraph_blocks(pdf_path, min_length=5):
#     """
#     Extrait le texte d’un PDF et le segmente en paragraphes, robustement.
#     Retourne une liste de dicts : [{text, page, index}]
#     """
#     import fitz  # PyMuPDF
#     blocks = []
#     doc = fitz.open(pdf_path)
#     for page_number, page in enumerate(doc, start=1):
#         text = page.get_text("text")
#         lines = text.splitlines()
#         paragraph = []
#         para_index = 1
#         for line in lines:
#             if line.strip() == "":
#                 # Fin de paragraphe
#                 if paragraph:
#                     para_text = " ".join(paragraph).strip()
#                     if len(para_text) >= min_length:
#                         blocks.append({
#                             "text": para_text,
#                             "page": page_number,
#                             "index": para_index
#                         })
#                         para_index += 1
#                     paragraph = []
#             else:
#                 paragraph.append(line.strip())
#         # Dernier paragraphe s'il y en a un
#         if paragraph:
#             para_text = " ".join(paragraph).strip()
#             if len(para_text) >= min_length:
#                 blocks.append({
#                     "text": para_text,
#                     "page": page_number,
#                     "index": para_index
#                 })
#     return blocks





# def highlight_diff_html(text1, text2):
#     """
#     Retourne un HTML où :
#     - les suppressions (text1) sont en rouge barré,
#     - les ajouts (text2) en vert souligné,
#     - les modifications sont surlignées.
#     """
#     d = difflib.ndiff(text1.split(), text2.split())
#     result = []

#     for token in d:
#         if token.startswith("- "):  # supprimé
#             result.append(f"<span style='color:red;text-decoration:line-through;'>{token[2:]}</span>")
#         elif token.startswith("+ "):  # ajouté
#             result.append(f"<span style='color:green;text-decoration:underline;'>{token[2:]}</span>")
#         elif token.startswith("? "):  # aide (non utilisée ici)
#             continue
#         else:
#             result.append(token[2:])

#     return "<p>" + " ".join(result) + "</p>"



##### AUTRE APPROCHE POUR EXTRAIRE LES BLOCS DE TEXTE:

def extract_paragraph_blocks(file_path, min_length=5):
    """
    Extrait le texte d'un PDF ou DOCX et le segmente en paragraphes robustement.
    Retourne une liste de dicts : [{text, page, index, type}]
    """
    file_path = Path(file_path)
    
    if file_path.suffix.lower() == '.pdf':
        return extract_pdf_paragraphs(file_path, min_length)
    elif file_path.suffix.lower() in ['.docx', '.doc']:
        return extract_docx_paragraphs(file_path, min_length)
    else:
        raise ValueError(f"Format de fichier non supporté: {file_path.suffix}")

def extract_pdf_paragraphs(pdf_path, min_length=5):
    """Extraction spécialisée pour PDF avec multiple techniques"""
    blocks = []
    doc = fitz.open(pdf_path)
    
    for page_number, page in enumerate(doc, start=1):
        # Méthode 1: Utiliser les blocs de texte PyMuPDF
        text_blocks = page.get_text("dict")["blocks"]
        
        para_index = 1
        for block in text_blocks:
            if "lines" in block:  # Bloc de texte
                # Extraire le texte de chaque bloc
                block_text = ""
                for line in block["lines"]:
                    for span in line["spans"]:
                        block_text += span["text"] + " "
                
                block_text = block_text.strip()
                if len(block_text) >= min_length:
                    # Diviser le bloc en paragraphes potentiels
                    paragraphs = split_text_into_paragraphs(block_text)
                    
                    for para_text in paragraphs:
                        if len(para_text.strip()) >= min_length:
                            blocks.append({
                                "text": para_text.strip(),
                                "page": page_number,
                                "index": para_index,
                                # "type": classify_text_type(para_text.strip())
                            })
                            para_index += 1
    
    doc.close()
    return blocks

def extract_docx_paragraphs(docx_path, min_length=5):
    """Extraction spécialisée pour DOCX"""
    blocks = []
    doc = Document(docx_path)
    
    para_index = 1
    current_page = 1  # Approximation pour DOCX
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if len(text) >= min_length:
            blocks.append({
                "text": text,
                "page": current_page,
                "index": para_index,
                # "type": classify_text_type(text),
                "style": paragraph.style.name if paragraph.style else "Normal"
            })
            para_index += 1
    
    return blocks


def split_text_into_paragraphs(text):
    """
    Divise un texte en paragraphes en utilisant plusieurs heuristiques
    """
    # Nettoyer le texte
    text = re.sub(r'\s+', ' ', text)  # Normaliser les espaces
    
    # Méthode 1: Diviser sur les doubles retours à la ligne
    paragraphs = re.split(r'\n\s*\n', text)
    
    # Si pas de division claire, utiliser d'autres heuristiques
    if len(paragraphs) == 1:
        # Méthode 2: Diviser sur les fins de phrase suivies d'espaces multiples
        paragraphs = re.split(r'(?<=[.!?])\s{2,}(?=[A-ZÀ-Ý])', text)
    
    # Si toujours pas de division, essayer les retours à la ligne simples
    if len(paragraphs) == 1:
        # Méthode 3: Diviser sur retours à la ligne + majuscule
        paragraphs = re.split(r'\n(?=[A-ZÀ-Ý])', text)
    
    # Méthode 4: Si le texte est très long, diviser par taille
    final_paragraphs = []
    for para in paragraphs:
        para = para.strip()
        if len(para) > 1000:  # Paragraphe très long, le diviser
            sentences = re.split(r'(?<=[.!?])\s+', para)
            current_chunk = ""
            for sentence in sentences:
                if len(current_chunk + sentence) < 500:
                    current_chunk += sentence + " "
                else:
                    if current_chunk:
                        final_paragraphs.append(current_chunk.strip())
                    current_chunk = sentence + " "
            if current_chunk:
                final_paragraphs.append(current_chunk.strip())
        else:
            if para:
                final_paragraphs.append(para)
    
    return final_paragraphs


